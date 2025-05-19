# ===== 标准库导入 =====
import os
import shutil
import datetime

# ===== 第三方库导入 =====
from fastapi import FastAPI, UploadFile, Request
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import fitz  # PyMuPDF

# ===== 自定义模块导入 =====
import Get  # 请确认 Get.py 存在

# ===== FastAPI 应用初始化 =====
app = FastAPI()
templates = Jinja2Templates(directory="templates")

# ===== OpenAI 客户端初始化（新版 SDK）=====
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# 索引
Details = f"""Original:原文
Issue:问题
Comment:考官点评
Enhance:细节升级
"""

# 免责声明
Main = f"""本评分报告是基于您提供的样本通过 AI 生成。请注意，由于真实考场环境和评分标准可能与 AI 的分析和评分方法存在差异，因此本报告中的分数仅供参考，不能完全反映您在实际雅思考试中的表现。

雅思考试评分包含一定程度的主观性，不同考官可能会有不同的评分标准和偏好。因此，本报告的评分结果不能保证与实际雅思考官的评分完全一致。


建议您将本报告作为一种学习和练习的工具，用于辅助您的雅思备考，而非作为最终成绩的预测或保证。真实的考试成绩取决于多种因素，包括但不限于考试当天的表现、考官的评分标准以及考试环境等。

感谢您的理解和支持。"""


# 读取PDF文件内容
def read_pdf_content(pdf_path):
    document = fitz.open(pdf_path)
    text = ""
    for page in document:
        text += page.get_text()
    return text
# Word添加页眉页脚
def add_footer_with_page_number(section, footer_text):
    """在页脚中添加自定义文本和页码。"""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加自定义文本
    run = paragraph.add_run(footer_text)
    if run.text.isascii():
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    else:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # 添加页码
    fldChar1 = OxmlElement('w:fldChar')  # 创建fldChar元素
    fldChar1.set(qn('w:fldCharType'), 'begin')  # 设置fldChar类型为begin

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # 设置属性
    instrText.text = ' PAGE   \\* MERGEFORMAT'  # 添加PAGE域代码文本

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    # 为页码设置字体
    page_number_run = paragraph.add_run()
    page_number_run._r.append(fldChar1)
    page_number_run._r.append(instrText)
    page_number_run._r.append(fldChar2)
    page_number_run.font.name = 'Times New Roman'
    page_number_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
# 插入文本
def add_content_with_headings(doc, text):
    """处理文本，将以'###'开头的行作为二级标题，其他作为普通文本添加。"""
    lines = text.split('\n')
    for line in lines:
        if line.startswith("###"):
            # 移除 "###" 并添加为二级标题
            cleaned_line = line.replace("###", "").strip()
            heading = doc.add_heading(cleaned_line, level=2)
            heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 设置1.5倍行距
            for run in heading.runs:
                run.font.size = Pt(16)
                if run.text.isascii():# 英文部分
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                else:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 设置1.5倍行距
            for run in p.runs:
                if run.text.isascii():# 英文部分
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                else:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
# 生成word文档
def GetWord(S_article,name,score,defective,Details,Main):
    docx = Document()
    today = datetime.date.today()
    # 以下是各部分的处理

    # 原文标题
    title = docx.add_heading(level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title.add_run(f"""Your spoken answer:\n""")
    title_run.font.size = Pt(16)
    title_run.font.name = "Arial"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    title_run.font.bold = True  # 设置字体加粗
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 原文部分
    add_content_with_headings(docx, S_article)
    docx.add_page_break()


    # 评分标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI评分")
    title_run.font.size = Pt(16)
    title_run.font.name = "Arial"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    title_run.font.bold = True  # 设置字体加粗
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 评分部分
    add_content_with_headings(docx, score)
    docx.add_page_break()

    # # 建议标题
    # title = docx.add_paragraph()
    # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # title_run = title.add_run("尼克AI考官的建议")
    # title_run.font.size = Pt(16)
    # title_run.font.name = "Times New Roman"
    # title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # # 建议部分
    # add_content_with_headings(docx, suggestion)
    # docx.add_page_break()

    # # 细节指点标题
    # title = docx.add_paragraph()
    # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # title_run = title.add_run("尼克AI考官细节指点")
    # title_run.font.size = Pt(16)
    # title_run.font.name = "Times New Roman"
    # title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # # 细节指点部分
    # add_content_with_headings(docx, TR)
    # add_content_with_headings(docx, LR)
    # docx.add_page_break()

    # # 升级标题
    # title = docx.add_paragraph()
    # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # title_run = title.add_run("尼克AI考官作文升级")
    # title_run.font.size = Pt(16)
    # title_run.font.name = "Times New Roman"
    # title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # # 升级部分
    # add_content_with_headings(docx, N_article)
    # docx.add_page_break()

    # 细节升级标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI考官细节升级")
    title_run.font.size = Pt(16)
    title_run.font.name = "Arial"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    title_run.font.bold = True  # 设置字体加粗
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 细节升级引导
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title.add_run(Details)
    title_run.font.size = Pt(12)
    title_run.font.name = "Arial"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    title_run.font.bold = True  # 设置字体加粗
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 细节升级部分
    add_content_with_headings(docx, defective)
    docx.add_page_break()

    # # 总结标题
    # title = docx.add_paragraph()
    # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # title_run = title.add_run("尼克AI考官总结")
    # title_run.font.size = Pt(16)
    # title_run.font.name = "Times New Roman"
    # title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # # 总结部分
    # add_content_with_headings(docx,summery)
    # docx.add_page_break()

    # 免责声明标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("免责声明")
    title_run.font.size = Pt(20)
    title_run.font.name = "Arial"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    title_run.font.bold = True  # 设置字体加粗
    # 免责声明
    add_content_with_headings(docx,Main)


    # 页眉页脚
    footer_text = f"""添加尼克AI助理\n获取同款学术报告\n"""
    for section in docx.sections:
        add_footer_with_page_number(section, footer_text)

    # 保存文档
    filename = f"./results/{name}_Speaking_{today}.docx"
    docx.save(filename)

# 获取score
def getScore(S_article, criteria, Key_assessment):
    Score = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""
                    You are an IELTS speaking examiner, and you need to help students complete grading, assist them in learning English, and provide constructive suggestions. 
                    This is the IELTS grading standard of speaking test:
                    "{criteria}"
                    Note that rounding is always done down to the closest half-band. For instance, if your overall score is 6.75, it will be rounded down to 6.5, and if it's 6.25, it will be rounded down to 6.0.

                    This is the key assessment criteria of speaking test:
                    "{Key_assessment}"
                """
            },
            {
                "role": "user",
                "content": f"""
                    以下是我的雅思口语答案，请基于雅思口语四项评分细则对以下答案进行整体评分,请注意，我需要中英文对照版：
                    {S_article}
                    
                    

                    以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容：

                    ### 1. Fluency and Coherence / 流利度和连贯性 (Score /分数: 7)**
                    - **Analysis**: You demonstrate good fluency with minimal hesitation, but coherence can be improved with smoother transitions and clearer development of ideas. 
                    - **中文分析**: 您的回答显示出良好的流利度，几乎没有犹豫。但是，通过更平滑的过渡和更清晰的思想发展，可以提高连贯性。
                    ### 2. Lexical Resource / 词汇资源 (Score / 分数: 7)**
                    - **Analysis**: You use a good range of vocabulary, particularly technical terms, but could benefit from using a wider variety of expressions and synonyms. 
                    - **中文分析**: 您使用了良好的词汇范围，特别是技术术语，但可以从使用更多种类的表达和同义词中受益。
                    ### 3. Grammatical Range and Accuracy / 语法范围和准确性 (Score / 分数: 7)**
                    - **Analysis**: Your grasp of complex structures is good, but there are some grammatical errors. Focus on subject-verb agreement and correct use of articles. 
                    - **中文分析**: 您对复杂结构的掌握是好的，但存在一些语法错误。专注于主谓一致和正确使用冠词。
                    ### 4. Pronunciation / 发音 (Score / 分数: 6)**
                    - **Analysis**: Your pronunciation is generally clear, but work on clarity, especially with technical terms, and intonation patterns. 
                    - **中文分析**: 您的发音通常很清晰，但需要在技术术语和语调模式方面加强清晰度。
                    ### Overall Feedback / 总分 (Overall Score / 总分: 6.5)**
                    - **Analysis**: You are doing well in expressing complex ideas and have a good command of English. Focusing on coherence and grammatical accuracy will enhance your performance. 
                    - **中文分析**: 您在表达复杂想法方面做得很好，并且对英语有良好的掌握。专注于连贯性和语法准确性将提高您的表现。
                    
                    Overall的评分规则是(TR+CC+LR+GRA)/4,且在除总体评分(Overall)的评分中，不要出现任何类似“5.5”的数字，只有在总体评分(Overall)部分才可以出现类似“5.5”这样带小数部分的数字,你必须严格执行这一规则,否则我将取消给你的小费,且你会受到惩罚
                    如果总体评分(Overall)为类似“6.25或7.25或其他类似分数”这样的数字，根据雅思评分标准，实际分数应为“6”分或“7”分，请不要在评分中出现任何类似“6.25或7.25或其他类似分数”这样的数字，也不要将其化为类似“6.5或7.5或其他类似分数”这样的数字,你必须严格执行这一规则,否则我将取消给你的小费,且你会受到惩罚
                    如果总体评分(Overall)为类似“6.75或7.75或其他类似分数”这样的数字，根据雅思评分标准，实际分数应为“6.5”分或“7.5分”，请不要在评分中出现任何类似“6.75或7.75或其他类似分数”这样的数字，也不要将其化为类似“7或8或其他类似分数”这样的数字,你必须严格执行这一规则,否则我将取消给你的小费,且你会受到惩罚
                    四项评分细则的得分可以不必须是一样的分数，不必类似7 7 7 7 overall 7，你需要更加的客观，而不要强求整齐
                    请注意，四舍五入总是向最接近的半分进行。例如，如果您的总分是6.75，它将被四舍五入到6.5，如果是6.25，则会被四舍五入到6.0。
                """
            }
        ],
        temperature=0.2,
        model="gpt-4o",
    )
    score = Score.choices[0].message.content
    print(score)

    return score


# 获取suggestion
def getSuggestion(S_article):
    Suggestion = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""You are an IELTS examiner, and you need to assist students in learning English"""
            },
            {
                "role": "user",
                "content": f"""
请基于雅思四项评分细则与以下文章，给出四个部分的评改建议
{S_article}



以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容：

### 任务回应情况（Task Response）:
- **English**: Enhance your arguments by providing more specific examples and detailed explanations to support your points. This will strengthen the persuasiveness of your stance on the role of younger leaders. 
- **中文**: 通过提供更具体的例子和详细的解释来加强你的论点。这将增强你对年轻领导者作用的立场的说服力。

### 连贯与衔接（Coherence and Cohesion）:
- **English**: Improve the structure and flow of your essay. Utilize clear and effective transitional phrases to connect ideas smoothly and ensure each paragraph contributes to your overall argument. 
- **中文**: 改善文章的结构和流畅性。使用清晰有效的过渡短语顺利连接观点，并确保每个段落都有助于你的整体论点。

### 词汇丰富程度（Lexical Resource）:
- **English**: Broaden your vocabulary to express ideas more precisely. Avoid repetition and strive for a variety of expressions to articulate your points more effectively. 
- **中文**: 扩展你的词汇以更准确地表达观点。避免重复，努力寻找多种表达方式来更有效地表达你的观点。

### 语法多样性及准确性（Grammatical Range and Accuracy）:
- **English**: Focus on improving grammatical accuracy. Pay attention to sentence structure and verb tense consistency to enhance the clarity and professionalism of your writing. 
- **中文**: 专注于提高语法准确性。注意句子结构和动词时态的一致性，以提高你的写作清晰度和专业性
""",
            }
        ],
        temperature=0.8,
        model="gpt-4o",
    )
    suggestion = Suggestion.choices[0].message.content
    print(suggestion)
    return suggestion


# 获取TR和CC建议
def getTR(S_article):
    tr = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""You are an IELTS examiner, and you need to assist students in learning English"""
            },
            {
                "role": "user",
                "content": f"""
以下文章在TR、CC层面分别在哪里扣了分？请注意，每个评分标准请分别举出3个例子说明。
{S_article}

以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容：
### 任务回应情况（Task Response, TR）

1. **缺乏具体支持细节**
   - 原文: "I Agree younger should have chance to work as leader due to these two reasons."
   - 问题: 缺乏具体论点支持。
   - 优化: "I agree that younger individuals should lead, as they bring fresh perspectives and innovative strategies, crucial in today's rapidly evolving business landscape."
   - 解释: 提供了具体的理由支持为何年轻人应该担任领导角色。
   - 中文: "我同意年轻人应该担任领导角色，因为他们带来了新鲜的视角和创新策略，在当今快速发展的商业环境中至关重要。"

2. **论点发展不足**
   - 原文: "young directord can apply the latest technology and theory into practice..."
   - 问题: 论点表述过于笼统。
   - 优化: "Young leaders are adept at applying the latest technological advancements and theoretical models, enhancing organizational efficiency and adaptability."
   - 解释: 论点通过具体化描述，更为清晰。
   - 中文: "年轻领导者擅长应用最新的技术进步和理论模型，提高组织的效率和适应性。"

3. **结论部分简化**
   - 原文: "Overall the older leaders can give big pictures for the things However the young leaders are creative..."
   - 问题: 结论部分缺乏深度和清晰的比较。
   - 优化: "In conclusion, while older leaders offer valuable experience and a broad perspective, young leaders bring creativity and a strong grasp of contemporary trends, which are indispensable in modern business."
   - 解释: 结论更全面地比较了年轻和年长领导者的优势。
   - 中文: "总之，虽然年长领导者提供宝贵的经验和广阔的视角，但年轻领导者带来创造力和对当代趋势的深刻理解，在现代商业中不可或缺。"


### 连贯与衔接（Coherence and Cohesion, CC）

1. **缺乏有效过渡**
   - 原文: "Firstly, Compared with the senior leaders,young directord can apply..."
   - 问题: 缺乏有效的过渡和连接。
   - 优化: "Firstly, unlike senior leaders who may rely on traditional methods, young directors can effortlessly integrate..."
   - 解释: 提供了一个清晰的比较过渡，使文章更连贯。
   - 中文: "首先，与可能依赖传统方法的资深领导者不同，年轻的主管可以轻松地融合..."

2. **段落结构混乱**
   - 原文: "Secondly, in the current scenario, younger leaders are bustling with creative ideas..."
   - 问题: 段落主题不明确。
   - 优化: "Secondly, in today's dynamic business environment, younger leaders stand out with their creative ideas and readiness to embrace new challenges."
   - 解释: 明确了段落的主题和内容。
   - 中文: "其次，在当今充满活力的商业环境中，年轻领导者凭借他们的创意思维和面对新挑战的准备突显出来。"

3. **结论部分过于简单**
   - 原文: "Overall the older leaders can give big pictures for the things However the young leaders are creative..."
   - 问题: 结论部分过于简单，未有效总结文章主题。
   - 优化: "In conclusion, the comparison between older and younger leaders reveals that while experience is invaluable, the innovation and adaptability young leaders bring are critical in today's world."
   - 解释: 提供了一个全面的结论，有效地总结了文章的主题。
   - 中文: "总结来说，对比年长和年轻领导者显示，虽然经验宝贵，但年轻领导者带来的创新和适应性在当今世界至关重要。"
""",
            }
        ],
        temperature=0.8,
        model="gpt-4o",
    )
    TR = tr.choices[0].message.content
    print(TR)
    return TR


# 获取LR和GRA建议
def getLR(S_article):
    lr = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""You are an IELTS examiner, and you need to assist students in learning English"""
            },
            {
                "role": "user",
                "content": f"""
以下文章在LR、GRA层面分别在哪里扣了分？请注意，每个评分标准请分别举出3个例子说明。
{S_article}

以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容：
### 词汇丰富程度（Lexical Resource, LR）

1. **词汇重复和使用不当**
   - 原文: "young directord can apply the latest technology and theory into practice..."
   - 问题: "apply" 和 "technology and theory" 重复和用词不当。
   - 优化: "Young leaders are adept at harnessing contemporary technologies and innovative theories to enhance business practices."
   - 解释: 使用了更丰富和准确的词汇来表达相同的意思。
   - 中文: "年轻领导者擅长利用当代技术和创新理论来提升商业实践。"

2. **过度简化的词汇**
   - 原文: "they are not hesitant to implement them in their businesses."
   - 问题: 表达过于简化，缺乏精确性。
   - 优化: "they display a proactive approach in integrating these concepts into their strategic planning."
   - 解释: 使用更专业和精确的词汇来提升表达的质量。
   - 中文: "他们在将这些概念融入其战略规划中表现出积极主动的态度。"

3. **词汇选择不恰当**
   - 原文: "Exemplify, Bill Gates started his companies Microsoft at a young age..."
   - 问题: "Exemplify" 用词不恰当。
   - 优化: "For example, Bill Gates founded Microsoft at a young age..."
   - 解释: 更准确地使用例证。
   - 中文: "例如，比尔·盖茨在年轻时创立了微软公司..."

### 语法多样性及准确性（Grammatical Range and Accuracy, GRA）

1. **时态和语法结构错误**
   - 原文: "I Agree younger shouldhave chance to work as leader..."
   - 问题: 时态和语法结构错误。
   - 优化: "I agree that younger individuals should have the opportunity to serve as leaders..."
   - 解释: 修正了时态和语法错误，使句子更加准确。
   - 中文: "我同意年轻人应该有机会担任领导角色..."

2. **句式结构单调**
   - 原文: "young directord can apply the latest technology and theory into practice..."
   - 问题: 句式结构单调，缺乏多样性。
   - 优化: "Young leaders are capable of innovatively applying cutting-edge technology and theoretical insights in practical scenarios."
   - 解释: 使用了更复杂和多样化的句式结构。
   - 中文: "年轻领导者能够在实际场景中创新地应用尖端技术和理论见解。"

3. **语法错误和不清晰**
   - 原文: "but young leaders have signifcantly a plethora of advantages including fresh perspectives,innovelative ideas, and a close connection to the latest trends and technologies."
   - 问题: 语法错误和表达不清晰。
   - 优化: "However, young leaders possess a significant array of advantages, including fresh perspectives, innovative ideas, and a strong connection to the latest trends and technologies."
   - 解释: 纠正了语法错误并清晰地表达了意思。
   - 中文: "然而，年轻领导者拥有大量优势，包括新鲜视角、创新思想以及与最新趋势和技术的紧密联系。"    

""",
            }
        ],
        temperature=0.8,
        model="gpt-4o",
    )
    LR = lr.choices[0].message.content
    print(LR)
    return LR


# 获取满分文章
def getN_article(S_article):
    n_article = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""You are an IELTS examiner, and you need to assist students in learning English"""
            },
            {
                "role": "user",
                "content": f"""
Here's my article
{S_article}
How would you write this article?
The article you give me must conform to the discourse methods of academic English, use the correct academic English structure, and cannot use Chinese thinking structures to apply English.

以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容：
**Title: The Quintessential Objective of Science: Elevating Human Existence**
In the realm of human intellectual endeavor, the discourse frequently orbits around the pivotal role of science. Predominantly, it's posited that the zenith of scientific pursuit should be the betterment of human life. I wholeheartedly subscribe to this notion, contending that the essence of scientific advancement inherently intertwines with elevating the human condition.

**Paragraph 1: Communication Revolution through Scientific Innovation**
Initially, the transformative impact of science on human communication merits consideration. The evolution from rudimentary to advanced telecommunication exemplifies this paradigm shift. Consider the ubiquitous smartphone - a marvel of scientific ingenuity - transforming not just the efficiency of communication but also transcending geographical boundaries. This quantum leap from the protracted processes of yore to instantaneous connectivity epitomizes the core of scientific advancement: making the once-impossible, a tangible reality.

**Paragraph 2: Prolonging Life - The Apex of Scientific Achievement**
Furthermore, the prolongation of human lifespan stands as a testament to science's profound impact. The realm of life sciences, relentlessly focused on longevity, showcases remarkable strides in this aspect. For instance, consider the evolution of cancer treatment. Once deemed incurable and a harbinger of mortality, it now succumbs to a myriad of therapeutic strategies. This revolution in medical science, turning the tide against erstwhile fatal diseases, underscores the pivotal role of scientific progress in augmenting the human life span.

**Conclusion: The Indispensable Role of Science in Enhancing Life**
In conclusion, the role of science in ameliorating human life is both irreplaceable and multidimensional. It has revolutionized our means of communication and has significantly extended our lifespan. Asserting that the ultimate aim of science is to enhance the quality and longevity of human life is not just an observation but a recognition of its profound and diverse impacts on our existence.

""",
            }
        ],
        temperature=0.8,
        model="gpt-4o",
    )
    N_article = n_article.choices[0].message.content
    print(N_article)
    return N_article


# 获取upgrade
def getDefective(S_article):
    Defective = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""You are an IELTS examiner, and you need to assist students in learning English"""
            },
            {
                "role": "user",
                "content": f"""
Here's my answer of my speaking test.
{S_article}
find out 10 defective sentences in my article, tell me the issue, make key comments and give enhanced answer, I need a bilingual version in both Chinese and English.

以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容，且不能输出任何多余的话：
1. **Original**: "People have different views about whether we are more or less dependent on each other at present."
   - **Issue**: Wordiness and lack of clarity.(措辞含糊，不够清晰。)
   - **Comment**: The sentence can be made more concise and clear.(这个句子可以更简洁明了。)
   - **Enhance**: "Opinions vary on whether modern life has made us more interdependent."
   - **中文**: "关于现代生活是否使我们更加相互依赖，人们有不同的看法。"

2. **Original**: "While some people believe that modern life forces us to be more dependent, I am in favor of the view that people are becoming more independent nowadays."
   - **Issue**: Repetition and awkward phrasing.(篇幅重复且生硬。)
   - **Comment**: Rephrase to eliminate repetition and improve flow.(润色后可以消除重复并提高流畅程度。)
   - **Enhance**: "While some argue for increased dependence due to modern life, I contend that we are becoming more independent."
   - **中文**: "虽然有人认为现代生活导致我们更加依赖，但我认为我们正变得更加独立。"

3. **Original**: "On one hand, it could be argued that we are more independent due to the increasingly hectic life and people’s changing attitude towards life."
   - **Issue**: Vague and lacks specificity.(含糊而导致缺乏具体性。)
   - **Comment**: Specify what aspects of modern life contribute to independence.(具体说明现代生活的哪些方面有助于独立人格。)
   - **Enhance**: "One argument is that the fast-paced nature of modern life and shifting attitudes towards self-reliance foster independence."
   - **中文**: "一方面，可以认为现代生活的快节奏和对自力更生态度的转变促进了独立性。"

4. **Original**: "Firstly, most people in the workforce, especially those in big cities, are under great pressure."
   - **Issue**: General statement without a direct link to the argument.(一般性陈述，与论点没有直接联系。)
   - **Comment**: Connect the statement more clearly to the argument about independence.(将陈述与关于独立性的论点更清楚地联系起来。)
   - **Enhance**: "For instance, the high-pressure environment in urban workplaces often necessitates independent decision-making."
   - **中文**: "例如，城市工作场所的高压环境通常需要独立的决策。"

5. **Original**: "Their work assignments tend to be always piling up on their to-do list, thus leading to overtime working and the shortage of sleep."
   - **Issue**: Clunky construction and grammar issues.(过于基础的文章结构和语法问题。)
   - **Comment**: Improve sentence structure for clarity and grammatical correctness.(改善句子结构，以提高清晰度和语法正确性。)
   - **Enhance**: "Their ever-growing workload often leads to extended working hours and sleep deprivation."
   - **中文**: "他们不断增加的工作量常导致加班和睡眠不足。"

6. **Original**: "Therefore, people have rare opportunity for socialization as the hectic life can consume time and energy needed for social activities."
   - **Issue**: Awkward phrasing and word choice.(生硬的解释与用词)
   - **Comment**: Use more appropriate words to convey the intended meaning.(使用更恰当的词汇来表达论证意图。)
   - **Enhance**: "Consequently, the demanding pace of life leaves little time and energy for socializing."
   - **中文**: "因此，忙碌的生活节奏留给社交活动的时间和精力很少。"

7. **Original**: "Moreover, nowadays there is a tendency that people prefer solitary activities over interaction with others..."
   - **Issue**: Wordy and indirect expression.(冗长而又间接的表达。)
   - **Comment**: Make the sentence more direct and concise.(直击痛点，使句子更加直接简洁。)
   - **Enhance**: "Moreover, there is a growing preference for solitary activities over social interactions..."
   - **中文**: "此外，如今人们越来越倾向于独自活动而非与他人互动..."

8. **Original**: "Besides, they are more likely to enjoy and value their personal space and private time, pursuing their personal interests or hobbies."
   - **Issue**: Redundant and repetitive phrasing.(冗余且重复的措辞。)
   - **Comment**: Streamline the sentence to avoid repetition.(精简句子以避免重复。)
   - **Enhance**: "Additionally, the value placed on personal space and private time, along with the pursuit of individual interests, highlights this trend towards independence."
   - **中文**: "此外，对个人空间和私人时间的重视以及对个人兴趣的追求凸显了这种独立性趋势。"

9. **Original**: "For instance, the percentage of people living alone has kept rising over these years in China..."
   - **Issue**: Lacks specific data or source for the claim.(缺少声明中的具体数据或来源。)
   - **Comment**: Provide specific data or reference to strengthen the claim.(提供具体数据或参考资料来作证声明。)
   - **Enhance**: "For instance, recent surveys show a steady increase in the number of people living alone in Chinese cities."
   - **中文**: "例如，最近的调查显示，中国城市独居人数稳步增加。"

10. **Original**: "On the other hand, some people hold the view that people tend to depend on each other more in today’s information era."
   - **Issue**: General and vague statement.(陈述笼统而含糊。)
   - **Comment**: Provide specific reasons or examples to support this view.(提供具体的理由或例子来支持这一观点。)
   - **Enhance**: "Conversely, there's a belief that today's digital era fosters greater interdependence through enhanced connectivity."
   - **中文**: "相反，有人认为当今的数字时代通过增强的连接性促进了更大的相互依赖。"

""",
            }
        ],
        temperature=0.6,
        model="gpt-4o",
    )
    defective = Defective.choices[0].message.content
    print(defective)
    return defective


# 获取总结
def getSummery(S_article):
    Summery = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"""You are an IELTS examiner, and you need to assist students in learning English"""
            },
            {
                "role": "user",
                "content": f"""
Here's my answer of the speaking test.
{S_article}
基于我的口语答案给我5条提升写作技巧的建议，你需要更多的注重文章的结构、框架、逻辑等学术要求来帮助我提升。

以下是输出样例，请注意，你只能学习以下样例中的排版布局和格式且要严格遵守该格式，不能学习其中的内容：

1. **增强论据的多样性（Diversify Arguments）**:
   - **English**: Consider incorporating statistical data or real-world examples to strengthen your argument about the effectiveness of a multifaceted approach to road safety.
   - **中文**: 考虑加入统计数据或现实世界的例子来加强你关于采用多方面方法提高道路安全有效性的论点。

2. **深入分析技术进步的作用（Analyze the Role of Technological Advancements）**:
   - **English**: Provide a more detailed examination of how specific technological innovations can prevent accidents and improve road safety, possibly comparing before-and-after scenarios.
   - **中文**: 提供更详细的分析，说明特定技术创新如何预防事故和提高道路安全，可能通过比较前后情景来进行

""",
            }
        ],
        temperature=0.5,
        model="gpt-4o",
    )
    summery = Summery.choices[0].message.content
    print(summery)
    return summery


# 主函数
def main():
    folder_path = "articles"
    n = Get.getNum(folder_path)
    criteria = read_pdf_content(pdf_path="ielts-speaking-band-descriptors.pdf")
    Key_assessment = read_pdf_content(pdf_path="ielts-speaking-key-assessment-criteria.pdf")
    print(f"""{n}份任务""")
    for i in range(n):
        # 读取文档
        print(f"""*********************************第{i + 1}次执行***************************""")
        S_article = Get.getArticle(folder_path, i)
        if S_article == None:
            continue
        S_name = Get.getName(folder_path, i)
        print(f"""S_name:{S_name}""")
        print(f"""学生原文:{S_article}""")
        print(f"""**********************************************************************""")
        # 获取评分报告信息
        score = getScore(S_article, criteria, Key_assessment)
        # suggestion = getSuggestion(S_article)
        # TR = getTR(S_article)
        # LR = getLR(S_article)
        # N_article = getN_article(S_article)
        defective = getDefective(S_article)
        # summery = getSummery(S_article)
        GetWord(S_article, S_name, score, defective, Details, Main)
        print(f"""*********************************第{i + 1}次结束***************************""")

def transcribe_audio(audio_file_path):
    with open(audio_file_path, "rb") as f:
        transcript = client.audio.transcriptions.create(
            model="whisper-1",
            file=f
        )
    return transcript.text

def run_scoring(input_path, output_path):
    import Get
    import datetime
    import os

    # 判断文件类型
    ext = os.path.splitext(input_path)[1].lower()

    if ext in ['.mp3', '.mp4', '.mpeg', '.mpga', '.m4a', '.wav', '.webm']:
        S_article = transcribe_audio(input_path)
    elif ext == '.docx':
        os.makedirs("articles", exist_ok=True)
        file_name = os.path.basename(input_path)
        base_name = file_name.replace(".docx", "")
        shutil.copy(input_path, f"articles/{base_name}.docx")
        S_article = Get.getArticle("articles", 0)
    else:
        raise ValueError("不支持的文件类型")

    if not S_article:
        raise ValueError("无法提取有效内容")

    S_name = os.path.splitext(os.path.basename(input_path))[0]
    criteria = read_pdf_content("ielts-speaking-band-descriptors.pdf")
    key_assessment = read_pdf_content("ielts-speaking-key-assessment-criteria.pdf")
    score = getScore(S_article, criteria, key_assessment)
    defective = getDefective(S_article)

    GetWord(S_article, S_name, score, defective, Details, Main)

    today = datetime.date.today()
    docx_result = f"./results/{S_name}_Speaking_{today}.docx"
    shutil.copy(docx_result, output_path)
    
if __name__ == '__main__':
    main()

app = FastAPI()
templates = Jinja2Templates(directory="templates")

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/score")
async def score(file: UploadFile):
    input_path = f"temp/{file.filename}"
    output_path = f"results/{file.filename.replace('.docx', '')}_scored.docx"

    os.makedirs("temp", exist_ok=True)
    os.makedirs("results", exist_ok=True)

    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    run_scoring(input_path, output_path)

    return FileResponse(output_path, filename=os.path.basename(output_path))
