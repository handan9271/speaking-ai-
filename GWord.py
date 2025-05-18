from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import results
import prompts
import time
import datetime



# name = input("\nPlease input your name:")
today = datetime.date.today()
name = prompts.S_name
# time.sleep(20)


# from docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_footer_with_page_number(section, footer_text):
    """在页脚中添加自定义文本和页码。"""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加自定义文本
    run = paragraph.add_run(footer_text)
    if run.text.isascii():
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加页码
    fldChar1 = OxmlElement('w:fldChar')  # 创建fldChar元素
    fldChar1.set(qn('w:fldCharType'), 'begin')  # 设置fldChar类型为begin

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # 设置属性
    instrText.text = ' PAGE   \\* MERGEFORMAT'  # 添加PAGE域代码文本

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    # 为页码设置字体
    page_number_run = paragraph.add_run()
    page_number_run._r.append(fldChar1)
    page_number_run._r.append(instrText)
    page_number_run._r.append(fldChar2)
    page_number_run.font.name = 'Times New Roman'
    page_number_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')



def add_content_with_headings(doc, text):
    """处理文本，将以'###'开头的行作为二级标题，其他作为普通文本添加。"""
    lines = text.split('\n')
    for line in lines:
        if line.startswith("###"):
            # 移除 "###" 并添加为二级标题
            cleaned_line = line.replace("###", "").strip()
            heading = doc.add_heading(cleaned_line, level=2)
            for run in heading.runs:
                run.font.size = Pt(16)
                if run.text.isascii():# 英文部分
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                if run.text.isascii():# 英文部分
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def GetWord():
    docx = Document()

    # 以下是各部分的处理

    # 原文标题
    title = docx.add_heading(level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title.add_run("您的原文:")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 原文部分
    add_content_with_headings(docx, results.article)
    docx.add_page_break()


    # 评分标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI评分")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 评分部分
    add_content_with_headings(docx, results.prompt1)
    docx.add_page_break()

    # 建议标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI考官的建议")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 建议部分
    add_content_with_headings(docx, results.prompt2)
    docx.add_page_break()

    # 细节指点标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI考官细节指点")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 细节指点部分
    add_content_with_headings(docx, results.prompt3)
    add_content_with_headings(docx, results.prompt4)
    docx.add_page_break()

    # 升级标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI考官作文升级")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 升级部分
    add_content_with_headings(docx, results.prompt6)
    docx.add_page_break()

    # 细节升级标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI考官细节升级")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 细节升级引导
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title.add_run(results.Details)
    title_run.font.size = Pt(12)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 细节升级部分
    add_content_with_headings(docx, results.prompt5)
    docx.add_page_break()

    # 总结标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("尼克AI考官总结")
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 总结部分
    add_content_with_headings(docx, results.prompt7)
    docx.add_page_break()

    # 免责声明标题
    title = docx.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("免责声明")
    title_run.font.size = Pt(20)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 免责声明
    add_content_with_headings(docx,results.Mian)


    # 页眉页脚
    footer_text = f"""添加尼克AI助理\n获取同款学术报告\n"""
    for section in docx.sections:
        add_footer_with_page_number(section, footer_text)

    # 保存文档
    filename = f"./test/{name}_{today}.docx"
    docx.save(filename)
